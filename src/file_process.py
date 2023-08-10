import re
from typing import Iterable, Dict
class FileProcess:
    @classmethod
    def passage_parse(cls, document, length=500, **kwargs) -> Dict:
        """
        passage parse
        :param document: document path
        :param length: passage length
        :return:xx
        """
        import os
        passage_dict = {}
        document_type = os.path.splitext(document)[-1]
        if document_type == ".docx":
            passage_dict.update(cls._docx_parse(document, length=1000))
        elif document_type == ".pdf":
            passage_dict.update(cls._pdf_parse(document, length=1000))
        elif document_type == ".txt":
            passage_dict.update(cls._txt_parse(document, length=1000))
        else:
            raise ValueError('''document %s type error, it's not .docx .pdf .txt''' % (document))

        # define result return
        result_dict = {
            'passages': passage_dict[document]
        }
        return result_dict
    
    @classmethod
    def split_text(cls, string: str, length) -> Iterable:
        """Split by length and end with specified punctuation"""
        res = []

        text_arr = re.split(r'([，；。！？ ,;!?])', string)
        text_arr.append("")
        text_arr = ["".join(i) for i in zip(text_arr[0::2], text_arr[1::2])]

        context = ""
        for index, text in enumerate(text_arr):
            context += text
            if len(context) == length:
                res.append(context)
            elif len(context) > length:
                for i in range(0, len(context), length):
                    res.append(context[i:i + length])
                context = ""
            elif index == len(text_arr) - 1:
                res.append(context.strip())

        return res

    @classmethod
    def _text_segment(cls, paragraphs, length, document):
        parse_res = {document: []}
        cur_passage = ''
        for index, p in enumerate(paragraphs):
            if len(p.strip()) == 0:
                continue
            text = p.strip().encode('utf-8', 'ignore').decode('utf-8')

            if len(text) > length:
                text_arr = cls.split_text(text, length)
                for t in text_arr:
                    if len(t.strip()) == 0:
                        continue
                    passage = t.strip()
                    if len(passage) + len(cur_passage) > length:
                        parse_res[document].append(cur_passage)
                        cur_passage = passage
                    else:
                        cur_passage += passage
            else:
                passage = text
                if len(passage) + len(cur_passage) > length:
                    parse_res[document].append(cur_passage)
                    cur_passage = passage
                else:
                    cur_passage += passage

        if cur_passage:
            parse_res[document].append(cur_passage)
        return parse_res

    @classmethod
    def _docx_parse(cls, document, length) -> Dict:
        """.docx parse"""

        from docx import Document

        # parse_res = {document: []}
        doc = Document(document)

        # for index, p in enumerate(doc.paragraphs):
        #     if len(p.text.strip()) == 0:
        #         continue
        #     if p.style.name.startswith("Normal"):
        #         text = p.text.strip().encode('utf-8', 'ignore').decode('utf-8')
        #         if len(text) > length:
        #             text_arr = cls.split_text(text, length)
        #             for t in text_arr:
        #                 t.strip()
        #                 if len(t.strip()) == 0:
        #                     continue
        #                 passage = t.strip()
        #                 parse_res[document].append(passage)
        #         else:
        #             passage = text
        #             parse_res[document].append(passage)
        # 小于length时，多段合到一起
        paragraphs = []
        for index, p in enumerate(doc.paragraphs):
            if len(p.text.strip()) == 0:
                continue
            if p.style.name.startswith("Normal"):
                paragraphs.append(p.text)
        parse_res = cls._text_segment(paragraphs, length, document)

        return parse_res

    @classmethod
    def _pdf_parse(cls, document, length) -> Dict:
        """.pdf parse"""

        import io
        from pdfminer.pdfpage import PDFPage
        from pdfminer.pdfinterp import PDFResourceManager
        from pdfminer.pdfinterp import PDFPageInterpreter
        from pdfminer.converter import TextConverter

        def extract_content_from_pdf(document):
            """
            This function extracts text from pdf file and return text as string.
            :param pdf_path: path to pdf file.
            :return: text string containing text of pdf.
            """
            resource_manager = PDFResourceManager(caching=False)
            fake_file_handle = io.StringIO(initial_value='\n', newline=None)
            converter = TextConverter(resource_manager, fake_file_handle, codec='utf-8', imagewriter=None)
            page_interpreter = PDFPageInterpreter(resource_manager, converter)

            with open(document, 'rb') as fh:
                for page in PDFPage.get_pages(fh, caching=True, check_extractable=True):
                    page_interpreter.process_page(page)

                text = fake_file_handle.getvalue()
            # close open handles
            converter.close()
            fake_file_handle.close()

            return text

        all_content = extract_content_from_pdf(document)
        passage_res = {document: []}

        # if all_content:
        #     for content in all_content.split("\f"):
        #         content = content.strip().encode('utf-8', 'ignore').decode('utf-8')
        #         if len(content) > length:
        #             text_arr = cls.split_text(content, length)
        #             for text in text_arr:
        #                 text = text.strip()
        #                 if len(text) == 0:
        #                     continue
        #                 passage = text
        #                 passage_res[document].append(passage)
        #         else:
        #             if len(content) == 0:
        #                 continue
        #             passage = content
        #             passage_res[document].append(passage)
        if all_content:
            passage_res = cls._text_segment(all_content.split("\f"), length, document)

        return passage_res


    @classmethod
    def _txt_parse(cls, document, length) -> Dict:
        """.txt parse"""

        f = open(document, "r", encoding="utf-8")
        passage_res = {document: []}

        # for line in f.readlines():
        #     line = line.strip().encode('utf-8', 'ignore').decode('utf-8')
        #     if len(line) > length:
        #         text_arr = cls.split_text(line, length)
        #         for text in text_arr:
        #             text = text.strip()
        #             if len(text) == 0:
        #                 continue
        #             passage = text
        #             passage_res[document].append(passage)
        #     else:
        #         if len(line) == 0:
        #             continue
        #         passage = line
        #         passage_res[document].append(passage)
        passage_res = cls._text_segment(f.readlines(), length, document)
        f.close()

        return passage_res

# d = FileProcess.passage_parse('price.pdf')
# print(d)